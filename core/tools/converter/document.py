# document.py 
import os
import re
from docx import Document
from PyPDF2 import PdfReader
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import arabic_reshaper
from bidi.algorithm import get_display

SUPPORTED_INPUT = ["docx", "pdf", "txt", "md"]
SUPPORTED_OUTPUT = ["pdf", "docx", "txt", "md"]
SUPPORTED_FORMATS = list(set(SUPPORTED_INPUT + SUPPORTED_OUTPUT))

BASE_DIR = os.path.dirname(__file__)
FONT_PATH = os.path.join(BASE_DIR, "DejaVuSans.ttf")

# ---------- TEXT UTILS ----------

def remove_emojis(text):
    return re.sub(r"[\U00010000-\U0010ffff]", "", text)

def prepare_rtl_text(text):
    fixed_lines = []
    for line in text.split("\n"):
        line = remove_emojis(line)
        reshaped = arabic_reshaper.reshape(line)
        bidi_text = get_display(reshaped)
        fixed_lines.append(bidi_text)
    return "\n".join(fixed_lines)

# ---------- CORE ----------

def extract_text(path, fmt):
    fmt = fmt.lower()
    text = ""

    if fmt == "txt" or fmt == "md":
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

    elif fmt == "docx":
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    elif fmt == "pdf":
        reader = PdfReader(path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

    return text.strip()

def write_pdf(text, output_path):
    if not os.path.exists(FONT_PATH):
        raise FileNotFoundError("DejaVuSans.ttf must be next to document.py")

    pdfmetrics.registerFont(TTFont("DejaVu", FONT_PATH))
    text = prepare_rtl_text(text)

    c = canvas.Canvas(output_path)
    c.setFont("DejaVu", 11)
    width, height = c._pagesize
    x_margin = 40
    y = height - 40
    line_height = 18

    for line in text.split("\n"):
        if y < 40:
            c.showPage()
            c.setFont("DejaVu", 11)
            y = height - 40
        c.drawRightString(width - x_margin, y, line)
        y -= line_height

    c.save()

def write_docx(text, output_path):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

def write_txt(text, output_path):
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)

def write_output(text, output_path, output_fmt):
    output_fmt = output_fmt.lower()
    if output_fmt == "pdf":
        write_pdf(text, output_path)
    elif output_fmt == "docx":
        write_docx(text, output_path)
    elif output_fmt == "txt" or output_fmt == "md":
        write_txt(text, output_path)
    else:
        raise ValueError(f"Unsupported output format: {output_fmt}")

# ---------- WRAPPER ----------

def convert(input_path, output_path, **kwargs):
    output_fmt = kwargs.get("output_fmt") or kwargs.get("fmt")
    if not output_fmt:
        raise ValueError("Missing output format")
    
    input_ext = os.path.splitext(input_path)[1][1:].lower()
    if input_ext not in SUPPORTED_INPUT:
        raise ValueError(f"Unsupported input format: {input_ext}")

    text = extract_text(input_path, input_ext)
    write_output(text, output_path, output_fmt)
