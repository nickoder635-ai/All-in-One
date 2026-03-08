import io
import os
import time
from ddgs import DDGS
import arabic_reshaper
from ollama import Client
from docx import Document
from dotenv import load_dotenv
from bidi.algorithm import get_display
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from contextlib import redirect_stdout, redirect_stderr
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet

ENV_PATH = ".env"

def validate_api_key(api_key: str, summarizer_model: str, translator_model: str) -> bool:
    if not summarizer_model or not translator_model:
        return False
    try:
        if not api_key:
            raise ValueError("API Key is empty. Set OLLAMA_API_KEY in .env first!")
            
        client = Client(
            host="https://ollama.com",
            headers={"Authorization": f"Bearer {api_key}"}
        )
        client.chat(model=summarizer_model, messages=[{"role": "user", "content": "ping"}])
        client.chat(model=translator_model, messages=[{"role": "user", "content": "ping"}])
        return True
    except Exception:
        return False

def run_research(
    topic: str,
    api_key: str,
    status_callback=None,  # callback برای UI
    summarizer_model: str = None,
    translator_model: str = None,
    target_language: str = "english",
    output_format: str = "txt",
    output_dir: str = "."
):
    if summarizer_model is None or translator_model is None:
        summarizer_model = os.getenv("OLLAMA_GPT_MODEL", "")
        translator_model = os.getenv("OLLAMA_MINIMAX_MODEL", "")

    client = Client(
        host="https://ollama.com",
        headers={"Authorization": f"Bearer {api_key}"}
    )

    language_display = target_language.capitalize()

    # ترجمه header topic
    header = topic
    if target_language.lower() != "english":
        try:
            response = client.chat(
                model=translator_model,
                messages=[
                    {"role": "system", "content": f"Translate this word to {language_display}."},
                    {"role": "user", "content": topic}
                ]
            )
            header = response["message"]["content"]
        except Exception:
            pass

    # ---------- SEARCH ----------
    self.progress.emit(f"🔍 Searching for '{self.topic}'...")
    results = []
    import io
    from contextlib import redirect_stdout, redirect_stderr

    f = io.StringIO()  # برای گرفتن خروجی stdout
    e = io.StringIO()  # برای گرفتن خروجی stderr

    try:
        with redirect_stdout(f), redirect_stderr(e):
            with DDGS() as ddgs:
                for r in ddgs.text(self.topic, max_results=3):
                    results.append(f"Title: {r['title']}\nURL: {r['href']}\nSnippet: {r['body']}\n")
    except Exception as ex:
        self.error.emit(f"Search Error: {str(ex)}")
        return

    search_text = "\n\n".join(results)
    self.progress.emit("✅ Search completed")
    time.sleep(1.5)

    # -------- Summarize ----------
    if status_callback:
        status_callback("🧩 Summarizing results...")
    response = client.chat(
        model=summarizer_model,
        messages=[
            {"role": "system", "content": "Summarize clearly as continuous prose."},
            {"role": "user", "content": search_text}
        ]
    )
    summary = response["message"]["content"]
    if status_callback:
        status_callback("✅ Summarization completed")
    time.sleep(1.5)

    # -------- Translate ----------
    translated = summary
    if target_language.lower() != "english":
        if status_callback:
            status_callback(f"🌐 Translating to {language_display}")
        response = client.chat(
            model=translator_model,
            messages=[
                {"role": "system", "content": f"Translate to fluent {language_display}."},
                {"role": "user", "content": summary}
            ]
        )
        translated = response["message"]["content"]
        if status_callback:
            status_callback("✅ Translation completed")

    # -------- Write Output ----------
    content = f"{header}\n\n{translated}"
    safe_topic = "".join(c if c.isalnum() or c in " _-" else "_" for c in topic)
    filename = f"{safe_topic} {language_display}.{output_format}"
    filepath = os.path.join(output_dir, filename)

    if output_format in ("txt", "md"):
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
    elif output_format == "docx":
        doc = Document()
        doc.add_paragraph(header)
        doc.add_paragraph("")  # خط خالی برای فاصله
        doc.add_paragraph(translated)
        doc.save(filepath)
    elif output_format == "pdf":
        font_path = "DejaVuSans.ttf"
        if not os.path.exists(font_path):
            raise Exception("DejaVuSans.ttf required in project root for PDF output.")

        pdfmetrics.registerFont(TTFont("DejaVu", font_path))
        doc = SimpleDocTemplate(filepath, pagesize=A4)
        elements = []
        style = ParagraphStyle(
            "Custom",
            parent=getSampleStyleSheet()["Normal"],
            fontName="DejaVu",
            fontSize=12
        )
        # هدر
        header_line = header
        if target_language.lower() in ("persian", "arabic", "turkish"):
            header_line = get_display(arabic_reshaper.reshape(header_line))
        elements.append(Paragraph(header_line, style))
        elements.append(Spacer(1, 12))  # فاصله بیشتر برای جداسازی

        # متن اصلی
        for line in translated.split("\n"):
            if target_language.lower() in ("persian", "arabic", "turkish"):
                line = get_display(arabic_reshaper.reshape(line))
            elements.append(Paragraph(line, style))
            elements.append(Spacer(1, 8))
        doc.build(elements)
    else:
        raise ValueError("Unsupported format.")

    if status_callback:
        status_callback(f"💾 Research Saved to: {filename}")

    return filepath